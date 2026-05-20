"""Remember pipeline routes — POST /api/v1/remember and task management."""
from __future__ import annotations

import hashlib
import logging
import os
import re
import shutil
import subprocess
import tempfile
import uuid
from datetime import datetime
from typing import Any, Dict, Optional

from flask import Blueprint, current_app, jsonify, make_response, request

from core.server.routes.helpers import (
    _get_processor,
    _get_queue,
    _parse_bool_query,
    _validate_graph_id,
    _validate_text_input,
    _validate_positive_int,
    _get_system_monitor,
    get_json_body,
    err,
    ok,
)
from core.server.routes._constants import _BOOL_TRUE, _BOOL_FALSE
from core.server.monitor import LOG_MODE_DETAIL
from core.server.task_queue import RememberTask
from core.llm.sanitize import sanitize_user_input

# Security: Maximum text length to prevent DoS
_MAX_TEXT_LENGTH = 10_000_000  # 10MB
_MAX_FILE_SIZE = 10_000_000  # 10MB
_ALLOWED_FILE_EXTENSIONS = {'.txt', '.text', '.md', '.markdown', '.json', '.html', '.htm', '.csv', '.log', '.pdf', '.docx', '.doc'}

logger = logging.getLogger(__name__)

remember_bp = Blueprint("remember", __name__)


# ── Request parsing helpers (module-level) ──────────────────────────────────

def _remember_get_str(name: str, post_json: Dict[str, Any]) -> str:
    if name in post_json and post_json[name] is not None:
        v = post_json[name]
        return (v if isinstance(v, str) else str(v)).strip()
    if request.method == "POST" and request.form and name in request.form:
        return (request.form.get(name) or "").strip()
    return (request.args.get(name) or "").strip()


def _parse_bool_value(v: Any) -> Optional[bool]:
    if isinstance(v, bool):
        return v
    if isinstance(v, int) and v in (0, 1):
        return bool(v)
    if isinstance(v, str):
        s = v.strip().lower()
        if s in _BOOL_TRUE:
            return True
        if s in _BOOL_FALSE:
            return False
    return None


def _remember_get_bool(name: str, post_json: Dict[str, Any]) -> Optional[bool]:
    if name in post_json:
        parsed = _parse_bool_value(post_json[name])
        if parsed is not None:
            return parsed
    if request.method == "POST" and request.form and name in request.form:
        parsed = _parse_bool_value(request.form.get(name))
        if parsed is not None:
            return parsed
    return _parse_bool_query(name)


# ── Extracted helpers ───────────────────────────────────────────────────────

def _parse_remember_input(post_json: Dict[str, Any]):
    """Parse and validate request input. Returns (text, source_name, load_cache, event_time) or error tuple."""
    text = _remember_get_str("text", post_json)

    # 如果 text 为空，尝试从 multipart 上传文件读取
    if not text and request.files:
        file = request.files.get("file")
        if file and file.filename:
            # Security: Validate file size
            file.seek(0, os.SEEK_END)
            file_size = file.tell()
            file.seek(0)
            if file_size > _MAX_FILE_SIZE:
                return err(f"文件大小超过限制 ({_MAX_FILE_SIZE / 1_000_000}MB)", 400)

            # Security: Validate file extension
            file_ext = os.path.splitext(file.filename)[1].lower()
            if file_ext and file_ext not in _ALLOWED_FILE_EXTENSIONS:
                return err(f"不支持的文件类型: {file_ext}", 400)

            content = file.read()
            try:
                text = _uploaded_file_to_markdown(content, file.filename, file_ext)
            except ValueError as exc:
                return err(str(exc), 400)
            except Exception as exc:
                logger.exception("Failed to convert uploaded file %s", file.filename)
                return err(f"文件转换失败: {exc}", 400)

    if not text:
        return err("缺少 text 或 file（必填其一）", 400)

    # Security: Validate text length
    if len(text) > _MAX_TEXT_LENGTH:
        return err(f"文本长度超过限制 ({_MAX_TEXT_LENGTH / 1_000_000}MB)", 400)

    # Security: Check for null bytes in text
    if '\x00' in text:
        return err("文本包含非法字符（null bytes）", 400)

    # Reject text with no alphanumeric/CJK content (punctuation-only like "..." or "!!!")
    import re as _re
    if not _re.search(r'[\w一-鿿぀-ゟ゠-ヿ가-힯]', text):
        return err("文本缺少有效内容（仅包含标点/空白字符）", 400)

    # Security: Sanitize for LLM prompt safety
    text, was_sanitized = sanitize_user_input(text)
    if was_sanitized:
        logger.warning("Remember input was sanitized due to security concerns")

    processor = _get_processor()
    sn = _remember_get_str("source_name", post_json)
    dn = _remember_get_str("doc_name", post_json)
    sd = _remember_get_str("source_document", post_json)
    # 如果从文件上传且未指定 source_name，用文件名
    if request.files and request.files.get("file") and request.files["file"].filename:
        if not sn and not dn and not sd:
            sn = request.files["file"].filename
    if sn or sd or dn:
        source_name = sn or sd or dn
    else:
        source_name = _generate_default_document_name(text, processor)

    load_cache = _remember_get_bool("load_cache_memory", post_json)
    if load_cache is None:
        # 任务入队时就固化默认值，避免服务重启或配置变更后语义漂移。
        load_cache = bool(getattr(processor, "load_cache_memory", False))

    # 以"首次接收请求的时间"为基准：若未传 event_time，则使用当前接收时间并持久化到 journal。
    receive_time = datetime.now()
    event_time: Optional[datetime] = receive_time
    et_str = _remember_get_str("event_time", post_json) or None
    if et_str:
        try:
            event_time = datetime.fromisoformat(et_str.replace("Z", "+00:00"))
        except ValueError:
            return err("event_time 需为 ISO 8601 格式", 400)

    return text, source_name, load_cache, event_time


def _generate_default_document_name(text: str, processor: Any) -> str:
    """Generate a short display document name for raw text remember requests."""
    title = ""
    llm_client = getattr(processor, "llm_client", None)
    if llm_client is not None and getattr(llm_client, "_endpoint_available", False):
        try:
            prompt = (
                "请为下面这段用户记忆生成一个中文 Markdown 文档名。\n"
                "要求：15个汉字以内；不要扩展名；不要引号；不要解释；保留核心主题。\n\n"
                f"内容：\n{text[:2000]}"
            )
            raw = llm_client._call_llm(
                prompt,
                system_prompt="你只输出一个短文档名。",
                max_retries=1,
                timeout=20,
                allow_mock_fallback=False,
                request_max_tokens_scale=0.05,
            )
            title = _sanitize_generated_title(raw)
        except Exception as exc:
            logger.debug("LLM document title generation failed: %s", exc)
    if not title:
        title = _fallback_document_title(text)
    return f"{title}.md"


def _uploaded_file_to_markdown(content: bytes, filename: str, file_ext: str) -> str:
    """Convert supported uploaded files into Markdown text for the remember pipeline."""
    ext = (file_ext or "").lower()
    display_name = os.path.basename(filename or "uploaded")
    if ext in {".md", ".markdown"}:
        return _decode_text_upload(content)
    if ext in {".txt", ".text", ".csv", ".log"}:
        return f"# {display_name}\n\n{_decode_text_upload(content)}"
    if ext == ".json":
        return f"# {display_name}\n\n```json\n{_decode_text_upload(content)}\n```"
    if ext in {".html", ".htm"}:
        return f"# {display_name}\n\n{_html_to_markdownish(_decode_text_upload(content))}"
    if ext == ".pdf":
        return f"# {display_name}\n\n{_pdf_to_markdown(content)}"
    if ext == ".docx":
        return f"# {display_name}\n\n{_docx_to_markdown(content)}"
    if ext == ".doc":
        return f"# {display_name}\n\n{_doc_to_markdown(content)}"
    raise ValueError(f"不支持的文件类型: {ext or '(无扩展名)'}")


def _decode_text_upload(content: bytes) -> str:
    encodings = []
    if content.startswith((b"\xff\xfe", b"\xfe\xff")):
        encodings.extend(["utf-16", "utf-16-le", "utf-16-be"])
    elif b"\x00" in content[:4096]:
        even_nulls = content[:4096:2].count(0)
        odd_nulls = content[1:4096:2].count(0)
        if odd_nulls > even_nulls:
            encodings.extend(["utf-16-le", "utf-16"])
        else:
            encodings.extend(["utf-16-be", "utf-16"])
    encodings.extend(["utf-8-sig", "utf-8", "gb18030", "cp936", "big5", "shift_jis"])
    detected = _detect_text_encoding(content)
    if detected:
        encodings.append(detected)
    encodings.append("latin-1")

    last_error = None
    for encoding in dict.fromkeys(encodings):
        try:
            text = content.decode(encoding)
            if "\x00" in text:
                last_error = ValueError("文本解码后仍包含 null bytes")
                continue
            return text
        except (UnicodeDecodeError, UnicodeError, ValueError) as exc:
            last_error = exc
            continue
    raise ValueError(f"文件编码错误，无法识别文本编码: {last_error}")


def _detect_text_encoding(content: bytes) -> str:
    try:
        from charset_normalizer import from_bytes  # type: ignore
        best = from_bytes(content).best()
        if best and best.encoding:
            return str(best.encoding)
    except Exception:
        pass
    try:
        import chardet  # type: ignore
        result = chardet.detect(content)
        if result and result.get("encoding") and float(result.get("confidence") or 0) >= 0.5:
            return str(result["encoding"])
    except Exception:
        pass
    return ""


def _html_to_markdownish(html: str) -> str:
    try:
        from bs4 import BeautifulSoup  # type: ignore
        soup = BeautifulSoup(html or "", "html.parser")
        for tag in soup(["script", "style"]):
            tag.decompose()
        return soup.get_text("\n")
    except Exception:
        text = re.sub(r"(?is)<(script|style).*?>.*?</\1>", "", html or "")
        text = re.sub(r"(?i)<br\s*/?>", "\n", text)
        text = re.sub(r"(?i)</p\s*>", "\n\n", text)
        text = re.sub(r"<[^>]+>", "", text)
        return text


def _pdf_to_markdown(content: bytes) -> str:
    errors = []
    with tempfile.NamedTemporaryFile(suffix=".pdf", delete=False) as tmp:
        tmp.write(content)
        tmp_path = tmp.name
    try:
        try:
            from pypdf import PdfReader  # type: ignore
            reader = PdfReader(tmp_path)
            pages = []
            for idx, page in enumerate(reader.pages, 1):
                pages.append(f"## Page {idx}\n\n{page.extract_text() or ''}".strip())
            text = "\n\n".join(pages).strip()
            if text:
                return text
        except Exception as exc:
            errors.append(f"pypdf: {exc}")
        try:
            import pdfplumber  # type: ignore
            with pdfplumber.open(tmp_path) as pdf:
                pages = [f"## Page {idx}\n\n{page.extract_text() or ''}".strip() for idx, page in enumerate(pdf.pages, 1)]
            text = "\n\n".join(pages).strip()
            if text:
                return text
        except Exception as exc:
            errors.append(f"pdfplumber: {exc}")
        try:
            import fitz  # type: ignore
            doc = fitz.open(tmp_path)
            pages = [f"## Page {idx}\n\n{page.get_text() or ''}".strip() for idx, page in enumerate(doc, 1)]
            doc.close()
            text = "\n\n".join(pages).strip()
            if text:
                return text
        except Exception as exc:
            errors.append(f"pymupdf: {exc}")
    finally:
        try:
            os.unlink(tmp_path)
        except OSError:
            pass
    raise ValueError("PDF 未能提取文本" + (f"（{'; '.join(errors[:2])}）" if errors else ""))


def _docx_to_markdown(content: bytes) -> str:
    try:
        from docx import Document  # type: ignore
    except Exception as exc:
        raise ValueError(f"DOCX 转换需要 python-docx: {exc}") from exc
    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
        tmp.write(content)
        tmp_path = tmp.name
    try:
        doc = Document(tmp_path)
        blocks = []
        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue
            style = (para.style.name or "").lower() if para.style else ""
            if "heading 1" in style:
                blocks.append(f"# {text}")
            elif "heading 2" in style:
                blocks.append(f"## {text}")
            elif "heading 3" in style:
                blocks.append(f"### {text}")
            else:
                blocks.append(text)
        for table in doc.tables:
            for row in table.rows:
                cells = [cell.text.strip().replace("\n", " ") for cell in row.cells]
                if any(cells):
                    blocks.append("| " + " | ".join(cells) + " |")
        text = "\n\n".join(blocks).strip()
        if not text:
            raise ValueError("DOCX 未提取到文本")
        return text
    finally:
        try:
            os.unlink(tmp_path)
        except OSError:
            pass


def _doc_to_markdown(content: bytes) -> str:
    antiword = shutil.which("antiword")
    if antiword:
        with tempfile.NamedTemporaryFile(suffix=".doc", delete=False) as tmp:
            tmp.write(content)
            tmp_path = tmp.name
        try:
            result = subprocess.run([antiword, tmp_path], capture_output=True, text=True, timeout=30)
            if result.returncode == 0 and result.stdout.strip():
                return result.stdout.strip()
        finally:
            try:
                os.unlink(tmp_path)
            except OSError:
                pass
    try:
        import win32com.client  # type: ignore
    except Exception as exc:
        raise ValueError("DOC 转换需要安装 Microsoft Word/pywin32 或 antiword；建议另存为 DOCX/PDF 后上传") from exc
    with tempfile.TemporaryDirectory() as td:
        doc_path = os.path.join(td, "input.doc")
        txt_path = os.path.join(td, "output.txt")
        with open(doc_path, "wb") as f:
            f.write(content)
        word = win32com.client.Dispatch("Word.Application")
        word.Visible = False
        try:
            doc = word.Documents.Open(doc_path)
            doc.SaveAs(txt_path, FileFormat=2)
            doc.Close(False)
            return _decode_text_upload(open(txt_path, "rb").read())
        finally:
            word.Quit()


def _sanitize_generated_title(raw: str) -> str:
    title = str(raw or "").strip()
    title = re.sub(r"```(?:json|text)?|```", "", title, flags=re.I).strip()
    if title.startswith("{"):
        try:
            import json as _json
            obj = _json.loads(title)
            if isinstance(obj, dict):
                title = str(obj.get("title") or obj.get("name") or obj.get("document_name") or title)
        except Exception:
            pass
    title = title.splitlines()[0].strip() if title else ""
    title = re.sub(r"^[\"'“”‘’《<]+|[\"'“”‘’》>]+$", "", title).strip()
    title = re.sub(r"[\s\t\r\n]+", "", title)
    title = re.sub(r"[\\/:*?\"<>|#\[\]{}]+", "", title)
    title = re.sub(r"\.(md|markdown|txt)$", "", title, flags=re.I)
    return title[:15] if title else ""


def _fallback_document_title(text: str) -> str:
    first_heading = re.search(r"^\s{0,3}#{1,6}\s+(.+?)\s*$", text or "", flags=re.M)
    if first_heading:
        title = _sanitize_generated_title(first_heading.group(1))
        if title:
            return title
    compact = re.sub(r"\s+", "", text or "")
    compact = re.sub(r"[\\/:*?\"<>|#\[\]{}]+", "", compact)
    compact = compact.strip("，。！？；：,.!?;:、 \t\r\n")
    return (compact[:15] or f"文本记忆{hashlib.sha256((text or '').encode('utf-8')).hexdigest()[:6]}")


def _build_remember_task(text: str, source_name: str, load_cache: bool,
                         event_time: Optional[datetime]) -> RememberTask:
    """Create a RememberTask with a generated task_id."""
    task_id = uuid.uuid4().hex
    return RememberTask(
        task_id=task_id,
        text=text,
        source_name=source_name,
        load_cache=load_cache,
        control_action=None,
        event_time=event_time,
        original_path="",
    )


def _handle_sync_wait(remember_queue, task_id: str, timeout: float):
    """Block until task completes or timeout. Returns Flask response."""
    done_task = remember_queue.wait_for_task(task_id, timeout=timeout)
    if done_task is None:
        return err(f"任务 {task_id} 未找到", 404)
    task_dict = remember_queue._task_to_dict(done_task)
    if done_task.status == "completed":
        return make_response(jsonify({
            "success": True,
            "data": {
                "task_id": task_id,
                "status": "completed",
                "result": done_task.result,
                **task_dict,
            },
        }), 200)
    elif done_task.status == "failed":
        return make_response(jsonify({
            "success": False,
            "data": {
                "task_id": task_id,
                "status": "failed",
                "error": done_task.error,
                **task_dict,
            },
        }), 500)
    else:
        # Timeout: still running, return current state with 202
        return make_response(jsonify({
            "success": True,
            "data": {
                "task_id": task_id,
                "status": done_task.status,
                "message": f"同步等待超时（{timeout}秒），任务仍在处理中。GET /api/v1/remember/tasks/{task_id} 继续轮询",
                **task_dict,
            },
        }), 202)


def _log_remember_request(text: str, source_name: str, event_time: Optional[datetime]):
    """Log the incoming remember request via system monitor or logger."""
    preview = (text[:80] + "…") if len(text) > 80 else text
    event_time_display = event_time.isoformat() if event_time else "未指定"
    system_monitor = _get_system_monitor()
    if system_monitor is not None:
        system_monitor.event_log.info(
            "Remember",
            f"收到({request.method}): source_name={source_name!r}, "
            f"文本长度={len(text)} 字符, event_time={event_time_display}"
        )
        if system_monitor.mode == LOG_MODE_DETAIL:
            system_monitor.event_log.info("Remember", f"内容预览: {preview!r}")
    else:
        logger.debug(
            "[Remember] 收到(%s): source_name=%r, 文本长度=%d 字符, event_time=%s",
            request.method, source_name, len(text), event_time_display,
        )


# ── POST /api/v1/remember ─────────────────────────────────────────────────

@remember_bp.route("/api/v1/remember", methods=["POST"])
def remember():
    """记忆写入：POST 请求发起异步任务，入队后立即返回 task_id。

    输入方式（三选一）：
      - JSON body 的 text 字段（适合短文本）
      - multipart/form-data 的 file 字段（适合长文本/文件上传）
      - JSON body 的 file_path 字段（仅限服务端本机文件）

    参数：
      - graph_id（可选）：目标图谱 ID，默认 "default"
      - text（可选）：正文
      - file（可选）：上传文件（multipart）
      - source_name / doc_name / source_document（可选）：来源名称，默认 api_input
      - load_cache_memory（可选）：
        true = 接续图谱中已有缓存链（同图任务需串行）
        false = 不接续外部缓存链，但任务内部滑窗仍续写自己的 cache 链（可并行）
      - event_time（可选）：ISO 8601 事件时间
      - wait（可选）：true 时同步等待完成再返回（默认 false，异步返回 202）
      - timeout（可选）：同步等待超时秒数（默认 300，仅 wait=true 时生效）

    返回：
      - wait=false（默认）：HTTP 202 + task_id（异步轮询模式）
      - wait=true：HTTP 200 + 完整结果（同步阻塞模式，适合 Agent 单次调用）
    """
    try:
        # Validate graph_id
        _validate_graph_id(request.graph_id)

        remember_queue = _get_queue()
        post_json: Dict[str, Any] = {}
        if request.method == "POST":
            pj = get_json_body()
            if isinstance(pj, dict):
                post_json = pj

        parsed = _parse_remember_input(post_json)
        # Check if it's an error response (tuple with Flask response)
        if isinstance(parsed, tuple) and len(parsed) == 2 and isinstance(parsed[1], int):
            return parsed  # error response
        # Otherwise it should be a 4-tuple with parsed values
        text, source_name, load_cache, event_time = parsed

        _log_remember_request(text, source_name, event_time)

        task = _build_remember_task(text, source_name, load_cache, event_time)
        remember_queue.submit(task)

        # Synchronous wait mode
        wait_mode = _remember_get_bool("wait", post_json)
        if wait_mode:
            timeout = 300
            timeout_str = _remember_get_str("timeout", post_json)
            if timeout_str:
                try:
                    timeout = _validate_positive_int(timeout_str, "timeout")
                    timeout = max(10, min(3600, timeout))
                except ValueError:
                    return err("timeout 必须为正整数", 400)
            return _handle_sync_wait(remember_queue, task.task_id, timeout)

        # Default async mode: return 202 immediately
        return make_response(jsonify({
            "success": True,
            "data": {
                "task_id": task.task_id,
                "task_seq": task.task_seq,
                "status": "queued",
                "message": f"已加入队列 (#{task.task_seq})。GET /api/v1/remember/tasks/{task.task_seq} 查询进度",
                "original_path": task.original_path,
            },
        }), 202)
    except ValueError as ve:
        return err(str(ve), 400)
    except Exception as e:
        return err(str(e), 500)


# ── GET/DELETE /api/v1/remember/tasks/<task_id> ───────────────────────────

@remember_bp.route("/api/v1/remember/tasks/<task_id>", methods=["GET", "DELETE"])
def remember_status(task_id: str):
    """查询或删除异步记忆写入任务；推荐使用 /api/v1/remember/tasks/<task_id>。"""
    try:
        remember_queue = _get_queue()
        if request.method == "DELETE":
            deleted, message, status = remember_queue.request_delete_task(task_id)
            if not deleted:
                if message == "任务不存在":
                    return err(f"任务不存在。GET /api/v1/remember/tasks 查看所有任务", 404)
                return err(message, 409)
            return ok({
                "task_id": task_id,
                "status": status,
                "message": message,
            })
        t = remember_queue.get_status(task_id)
        if t is None:
            return err(f"任务不存在。GET /api/v1/remember/tasks 查看所有任务", 404)
        data: Dict[str, Any] = remember_queue._task_to_dict(t)
        data["original_path"] = t.original_path
        if t.status == "completed" and t.result:
            data["result"] = t.result
        if t.status == "failed" and t.error:
            data["error"] = t.error
        return ok(data)
    except Exception as e:
        return err(str(e), 500)


# ── POST /api/v1/remember/tasks/<task_id>/pause ───────────────────────────

@remember_bp.route("/api/v1/remember/tasks/<task_id>/pause", methods=["POST"])
def remember_pause(task_id: str):
    try:
        remember_queue = _get_queue()
        ok_pause, message, status = remember_queue.request_pause_task(task_id)
        if not ok_pause:
            if message == "任务不存在":
                return err(message, 404)
            return err(message, 409)
        return ok({
            "task_id": task_id,
            "status": status,
            "message": message,
        })
    except Exception as e:
        return err(str(e), 500)


# ── POST /api/v1/remember/tasks/<task_id>/resume ──────────────────────────

@remember_bp.route("/api/v1/remember/tasks/<task_id>/resume", methods=["POST"])
def remember_resume(task_id: str):
    try:
        remember_queue = _get_queue()
        ok_resume, message, status = remember_queue.resume_task(task_id)
        if not ok_resume:
            if message == "任务不存在":
                return err(message, 404)
            return err(message, 409)
        return ok({
            "task_id": task_id,
            "status": status,
            "message": message,
        })
    except Exception as e:
        return err(str(e), 500)


# ── GET /api/v1/remember/tasks ────────────────────────────────────────────

@remember_bp.route("/api/v1/remember/tasks", methods=["GET"])
def remember_queue_list():
    """查看记忆写入任务队列；推荐使用 /api/v1/remember/tasks。"""
    try:
        remember_queue = _get_queue()
        limit = min(request.args.get("limit", 50, type=int), 200)
        tasks = remember_queue.list_tasks(limit=limit)
        return ok({"tasks": tasks, "count": len(tasks)})
    except Exception as e:
        return err(str(e), 500)


# ── GET /api/v1/remember/monitor ──────────────────────────────────────────

@remember_bp.route("/api/v1/remember/monitor", methods=["GET"])
def remember_monitor():
    """返回 remember 的实时监控快照，适合 watch 或外部面板轮询。"""
    try:
        system_monitor = _get_system_monitor()
        detail = system_monitor.graph_detail(request.graph_id) if system_monitor else None
        if detail is None:
            remember_queue = _get_queue()
            limit = request.args.get("limit", 6, type=int)
            return ok({
                "graph_id": request.graph_id,
                "queue": remember_queue.get_monitor_snapshot(limit=limit),
            })
        return ok({
            "graph_id": request.graph_id,
            "storage": detail["storage"],
            "queue": detail["queue"],
            "threads": detail["threads"],
            "pipeline": detail.get("pipeline"),
        })
    except Exception as e:
        return err(str(e), 500)

